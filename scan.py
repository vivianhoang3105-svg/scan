import streamlit as st
from PIL import Image, ImageOps
import pytesseract
import io
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document

# --- 1. CẤU HÌNH TRANG ---
st.set_page_config(page_title="Máy Quét VIP Linh Linh", page_icon="🌿⚔️")
st.write("Fighting")
# --- 2. GIAO DIỆN VÀ MÀU SẮC (CSS) ---
st.markdown("""
    <style>
    /* Nền ứng dụng màu Vàng Kem */
    .stApp { background-color: #FEFDF5; }
    
    /* Tiêu đề chính */
    h1 { color: #2E7D32 !important; text-shadow: 1px 1px 2px white; font-family: 'Segoe UI', sans-serif;}
    
    /* FIX LỖI TÀNG HÌNH: Ép chữ của Nút Gạt và Kéo thả thành Xanh Lá Cây đậm */
    p, label, .stMarkdown p { color: #2E7D32 !important; font-weight: 500; font-size: 16px; }

    /* Nút Kích Hoạt Xanh Lá */
    div.stButton > button { 
        background-color: #2E7D32 !important; 
        border: 2px solid #A5D6A7 !important; 
        border-radius: 12px; 
    }
    /* Sơn Nút Gạt lúc TẮT (OFF) thành Xanh Lá Nhạt cho dễ nhìn */
    div[data-testid="stWidgetToggle"] label div[data-checked="false"] {
        background-color: #C8E6C9 !important; 
    }
    /* Chữ bên trong Nút Kích Hoạt là MÀU TRẮNG */
    div.stButton > button p { color: white !important; font-weight: bold; font-size: 16px !important; }
    
    /* Khung text hiển thị văn bản */
    .stTextArea textarea { 
        background-color: #FFFDF0 !important; 
        color: #1A531A !important; 
        border: 2px solid #2E7D32 !important; 
        font-size: 17px !important; 
        font-family: 'Segoe UI', sans-serif;
    }
    .stTextArea label p { color: #2E7D32 !important; font-weight: bold; font-size: 18px !important; }
    
    /* Nút Tải về Màu Cam */
    div.stDownloadButton > button { 
        background-color: #E65100 !important; 
        border: 2px solid #FFD700 !important; 
    }
    /* Chữ bên trong Nút Tải về là MÀU TRẮNG */
    div.stDownloadButton > button p { color: white !important; font-weight: bold; font-size: 18px !important; }
    
    /* Sơn Nút Gạt (Toggle) thành Xanh Lá khi được bật */
    div[data-testid="stWidgetToggle"] label div[data-checked="true"] {
        background-color: #2E7D32 !important;
    }
    
    footer {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)

# --- 3. BỘ NÃO AI GEMINI ---
def get_ai_response(content):
    api_key = st.secrets.get("GOOGLE_API_KEY")
    if not api_key: return None, "Chưa dán API Key vô Secrets"
    genai.configure(api_key=api_key)
    try:
        target_model = None
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                if 'gemini' in m.name.lower() and 'vision' not in m.name.lower():
                    target_model = m.name
                    break
        if not target_model: return None, "Không tìm thấy model AI."
        model = genai.GenerativeModel(target_model)
        prompt = f"Hãy sửa hết lỗi chính tả và OCR trong văn bản sau. Chỉ trả về kết quả đã sửa sạch đẹp, ngay ngắn:\n\n{content}"
        response = model.generate_content(prompt)
        return response.text, None
    except Exception as e: return None, f"Lỗi AI: {str(e)}"

# --- 4. HÀM TẠO FILE WORD ---
def create_word_file(text_content, title="VĂN BẢN TRÍCH XUẤT"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text_content)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.alignment = 0
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 5. HIỂN THỊ TRÊN WEB ---
st.title("🌿⚔️ Máy Quét VIP Linh Linh")
st.write("Vũ khí số hóa văn bản của Nữ hoàng Thiên Phủ - An toàn tuyệt đối 100%!")

uploaded_files = st.file_uploader("Kéo thả Ảnh hoặc PDF vào đây", type=["jpg", "jpeg", "png", "pdf"], accept_multiple_files=True)

# CÔNG TẮC BẢO MẬT (Mặc định là tắt)
use_ai = st.toggle("✨ Bật chế độ AI sửa lỗi (Lưu ý: Dữ liệu sẽ được gửi qua Google AI để xử lý)", value=False)

if uploaded_files:
    if st.button("Kích Hoạt Máy Quét 🚀"):
        all_raw_text = ""
        images_to_show = []

        with st.spinner("Đang soi từng nét chữ, đợi xíu nha..."):
            try:
                for file in uploaded_files:
                    if file.name.lower().endswith('.pdf'):
                        doc = fitz.open(stream=file.read(), filetype="pdf")
                        for page_index in range(len(doc)):
                            page = doc.load_page(page_index)
                            pix = page.get_pixmap()
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            images_to_show.append((img, f"PDF: {file.name} (T{page_index+1})"))
                            all_raw_text += f"\n{pytesseract.image_to_string(img, lang='vie+eng')}\n"
                    else:
                        img = ImageOps.exif_transpose(Image.open(file)).convert('RGB')
                        images_to_show.append((img, f"Ảnh: {file.name}"))
                        all_raw_text += f"\n{pytesseract.image_to_string(img, lang='vie+eng')}\n"

                # Hiện ảnh đã tải lên
                for pic, cap in images_to_show:
                    st.image(pic, caption=cap, width=500)

                # PHẦN 1: LUÔN HIỆN BẢN THÔ (AN TOÀN NỘI BỘ)
                st.text_area("📄 Văn Bản Thô (Bảo mật 100% nội bộ):", all_raw_text, height=250)
                raw_word_bytes = create_word_file(all_raw_text, "VĂN BẢN THÔ - BẢO MẬT")
                st.download_button(
                    label="📥 TẢI XUỐNG BẢN THÔ (.docx)",
                    data=raw_word_bytes,
                    file_name="VanBan_Tho_BaoMat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # PHẦN 2: NẾU BẬT CÔNG TẮC AI MỚI CHẠY ĐOẠN NÀY
                if use_ai and all_raw_text.strip():
                    st.divider() # Kẻ 1 đường ngang phân cách
                    with st.spinner("Đang kết nối AI Gemini để dọn dẹp chính tả..."):
                        corrected, error = get_ai_response(all_raw_text)
                        if corrected:
                            st.success("✨ AI ĐÃ SỬA XONG MƯỢT MÀ!")
                            st.text_area("💎 VĂN BẢN HOÀN HẢO (Đã qua AI):", corrected, height=450)
                            
                            ai_word_bytes = create_word_file(corrected, "VĂN BẢN ĐÃ QUA AI CHỈNH SỬA")
                            st.download_button(
                                label="📥 TẢI XUỐNG BẢN ĐÃ QUA AI (.docx)",
                                data=ai_word_bytes,
                                file_name="VanBan_AI_HoanHao.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        else:
                            st.warning(f"AI đang bận: {error}")
                            
            except Exception as e:
                st.error(f"Lỗi hệ thống: {e}")
  
